import pdfplumber
from langchain_core.documents import Document
from docx import Document as DocxDocument
import re

def enrich_pdf_chunks(pdf_path: str) -> list:
    chunks = []

    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            # Extract tables first as structured text
            table_texts = []
            for table in page.extract_tables():
                if not table:
                    continue
                rows = []
                for row in table:
                    cleaned = [cell.strip() if cell else "" for cell in row]
                    non_empty = [c for c in cleaned if c]
                    if non_empty:
                        rows.append(" | ".join(non_empty))
                if rows:
                    table_texts.append("\n".join(rows))

            # Extract text words with font size for heading detection
            words = page.extract_words(
                extra_attrs=["size", "fontname"],
                keep_blank_chars=False
            )
            if not words:
                continue

            # Detect dominant body font size
            sizes = [w.get("size", 12) for w in words]
            body_size = sorted(sizes)[len(sizes) // 2]  # median
            heading_threshold = body_size * 1.15

            # Rebuild text with heading markers
            lines = []
            current_line = []
            current_y = None

            for word in words:
                y = round(word["top"], 1)
                if current_y is None:
                    current_y = y
                if abs(y - current_y) > 3:
                    if current_line:
                        line_text = " ".join(current_line)
                        lines.append(line_text)
                    current_line = [word["text"]]
                    current_y = y
                else:
                    current_line.append(word["text"])
            if current_line:
                lines.append(" ".join(current_line))

            # Split into sections by heading detection
            current_heading = f"Page {page_num + 1}"
            current_body = []

            for line in lines:
                line_words = [w for w in words
                              if line.startswith(w["text"][:min(5, len(w["text"]))])]
                is_heading = False
                if line_words:
                    avg_size = sum(
                        w.get("size", body_size) for w in line_words
                    ) / len(line_words)
                    is_heading = (
                        avg_size >= heading_threshold and
                        len(line) < 120 and
                        len(line.split()) <= 12
                    )

                if is_heading and current_body:
                    body_text = "\n".join(current_body)
                    if table_texts:
                        body_text += "\n\n" + "\n\n".join(table_texts)
                        table_texts = []
                    enriched = (
                        f"SECTION: {current_heading}\n"
                        f"Keywords: policy, procedures, guidelines, onboarding, "
                        f"processes, workflows, documentation, organization.\n\n"
                        f"{body_text}"
                    )
                    chunks.append(Document(
                        page_content=enriched,
                        metadata={
                            "source": f"document_page_{page_num + 1}",
                            "section_title": current_heading
                        }
                    ))
                    current_heading = line
                    current_body = []
                else:
                    current_body.append(line)

            # Flush remaining content
            if current_body:
                body_text = "\n".join(current_body)
                if table_texts:
                    body_text += "\n\n" + "\n\n".join(table_texts)
                enriched = (
                    f"SECTION: {current_heading}\n"
                    f"Keywords: policy, procedures, guidelines, onboarding, "
                    f"processes, workflows, documentation, organization.\n\n"
                    f"{body_text}"
                )
                chunks.append(Document(
                    page_content=enriched,
                    metadata={
                        "source": f"document_page_{page_num + 1}",
                        "section_title": current_heading
                    }
                ))

    return chunks if chunks else [Document(
        page_content="",
        metadata={"source": "document_page_1"}
    )]

def chunk_docx_with_metadata(docx_path: str) -> list:
    doc = DocxDocument(docx_path)

    HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3"}
    MIN_CHUNK_SIZE = 100

    current_heading = "Introduction"
    current_paragraphs = []
    sections = []  # list of (heading, paragraphs_text)

    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    def _is_bold_heading(para):
        """Detect bold paragraphs used as visual headings (no heading style)."""
        text = para.text.strip()
        if not text or len(text) > 80:
            return False
        runs = [r for r in para.runs if r.text.strip()]
        if not runs:
            return False
        return all(r.bold for r in runs)

    def _extract_table_text(table):
        """Extract table as row-oriented text preserving column relationships."""
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # Deduplicate merged cells (python-docx repeats merged cell text)
            deduped = []
            for c in cells:
                if c and (not deduped or c != deduped[-1]):
                    deduped.append(c)
            if deduped:
                rows_text.append(" | ".join(deduped))
        return "\n".join(rows_text)

    for child in doc.element.body:
        if child in para_map:
            para = para_map[child]
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style and para.style.name else ""
            if style_name in HEADING_STYLES:
                sections.append((current_heading, "\n".join(current_paragraphs)))
                current_heading = text
                current_paragraphs = []
            elif _is_bold_heading(para):
                sections.append((current_heading, "\n".join(current_paragraphs)))
                current_heading = text
                current_paragraphs = []
            else:
                current_paragraphs.append(text)
        elif child in table_map:
            table = table_map[child]
            current_paragraphs.append(_extract_table_text(table))

    sections.append((current_heading, "\n".join(current_paragraphs)))

    chunks = []
    for heading, body in sections:
        if not body.strip():
            continue

        enriched_text = (
            f"SECTION: {heading}\n"
            f"Keywords: policy, procedures, guidelines, onboarding, processes, workflows, documentation, organization.\n\n"
            f"{body}"
        )

        if len(body) < MIN_CHUNK_SIZE and chunks:
            chunks[-1] = Document(
                page_content=chunks[-1].page_content + "\n\n" + enriched_text,
                metadata=chunks[-1].metadata
            )
        else:
            chunks.append(Document(
                page_content=enriched_text,
                metadata={"source": "orientation_guide", "section_title": heading}
            ))

    return chunks